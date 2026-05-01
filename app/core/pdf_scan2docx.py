#  Copyright (c) 2017-2026 null. All rights reserved.
"""PDF扫描件转Word模块，通过OCR技术识别图片中的文字并生成Word文档"""
import fitz
from docx import Document
from rapidocr_onnxruntime import RapidOCR

from app.utils.file_utils import get_file_path

DEFAULT_OCR_DPI = 200

_ocr_engine = None


def _get_ocr_engine():
    """
    获取OCR引擎单例

    :return: RapidOCR 引擎实例
    """
    global _ocr_engine
    if _ocr_engine is None:
        _ocr_engine = RapidOCR()
    return _ocr_engine


def _pdf_page_to_image(page, dpi: int = DEFAULT_OCR_DPI) -> bytes:
    """
    将PDF单页转换为PNG图片字节

    :param page: fitz.Page 对象
    :param dpi: 图片DPI，值越大识别精度越高但速度越慢
    :return: PNG格式的图片字节数据
    """
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat)
    return pix.tobytes("png")


def _ocr_image(image_bytes: bytes) -> str:
    """
    对图片执行OCR识别

    :param image_bytes: 图片字节数据
    :return: 识别出的文字内容，多行以换行符分隔
    """
    ocr = _get_ocr_engine()
    result, _ = ocr(image_bytes)
    if not result:
        return ""
    lines = []
    for item in result:
        lines.append(item[1])
    return "\n".join(lines)


def convert_pdf_scan_to_word(pdf_filename: str, word_filename: str, dpi: int = DEFAULT_OCR_DPI) -> str:
    """
    PDF扫描件（图片）转Word，通过OCR识别图片中的文字

    :param pdf_filename: 上传的PDF文件名（仅文件名，不含路径）
    :param word_filename: 输出的Word文件名（仅文件名，不含路径）
    :param dpi: PDF转图片的DPI，值越大识别精度越高但速度越慢
    :return: Word文件的完整路径
    :raises RuntimeError: 转换过程失败时抛出
    """
    pdf_path = get_file_path(pdf_filename)
    word_path = get_file_path(word_filename)

    try:
        doc = Document()
        with fitz.open(pdf_path) as pdf_doc:
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                image_bytes = _pdf_page_to_image(page, dpi=dpi)
                text = _ocr_image(image_bytes)

                if page_num > 0:
                    doc.add_page_break()

                heading = doc.add_heading(f"第 {page_num + 1} 页", level=2)
                heading.alignment = 0

                if text.strip():
                    for line in text.split("\n"):
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph("（此页未识别到文字内容）")

        doc.save(word_path)
        return word_path
    except Exception as e:
        raise RuntimeError(f"PDF扫描件转Word失败：{str(e)}")
